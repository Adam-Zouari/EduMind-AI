"""
DOCX extraction using python-docx
"""
from docx import Document
from pathlib import Path
from typing import Dict
import time
from core.base_extractor import BaseExtractor, ExtractionResult

class DOCXExtractor(BaseExtractor):
    """Extract text and metadata from DOCX files"""
    
    def extract(self, file_path: Path, **kwargs) -> ExtractionResult:
        """Extract text from DOCX"""
        start_time = time.time()
        self.logger.info(f"Extracting DOCX: {file_path}")
        
        try:
            doc = Document(file_path)
            
            # Extract paragraphs
            text_parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            # Extract tables
            tables_text = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text.strip() for cell in row.cells])
                    if row_text.strip():
                        tables_text.append(row_text)
            
            # Combine text
            text = "\n\n".join(text_parts)
            if tables_text:
                text += "\n\n--- TABLES ---\n\n" + "\n".join(tables_text)
            
            # Extract metadata
            core_props = doc.core_properties
            metadata = {
                "title": core_props.title or "",
                "author": core_props.author or "",
                "subject": core_props.subject or "",
                "keywords": core_props.keywords or "",
                "created": str(core_props.created) if core_props.created else "",
                "modified": str(core_props.modified) if core_props.modified else "",
                "num_paragraphs": len(doc.paragraphs),
                "num_tables": len(doc.tables),
                "extractor": "python-docx"
            }
            
            extraction_time = time.time() - start_time
            
            return ExtractionResult(
                text=text,
                metadata=metadata,
                format_type="docx",
                file_path=str(file_path),
                extraction_time=extraction_time,
                success=True
            )
            
        except Exception as e:
            self.logger.error(f"DOCX extraction failed: {e}")
            return self._create_error_result(file_path, str(e))